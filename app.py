#!/usr/bin/env python3
"""
PDF文档解析查看器 - Flask Web应用
"""
import os
import uuid
import threading
import json
import base64
import shutil
import time
import requests
from requests.adapters import HTTPAdapter
from datetime import datetime
from pathlib import Path
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor, as_completed

from flask import Flask, request, jsonify, send_file, send_from_directory, render_template, after_this_request
from flask_cors import CORS

from document_parser import DocumentParser
from config import DEEPLX_API_URL, DEFAULT_SOURCE_LANG, DEFAULT_TARGET_LANG

app = Flask(__name__)
CORS(app)

BASE_DIR = Path(__file__).resolve().parent
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'outputs')
FRONTEND_DIST = BASE_DIR / 'frontend' / 'dist'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

tasks = {}
DEEPLX_MAX_CONCURRENCY = int(os.environ.get("DEEPLX_MAX_CONCURRENCY", "4"))
DEEPLX_TIMEOUT = float(os.environ.get("DEEPLX_TIMEOUT", "30"))
DEEPLX_MAX_RETRIES = int(os.environ.get("DEEPLX_MAX_RETRIES", "2"))
DEEPLX_RATE_LIMIT = float(os.environ.get("DEEPLX_RATE_LIMIT", "0.3"))
DEEPLX_HEALTH_TTL = float(os.environ.get("DEEPLX_HEALTH_TTL", "60"))
DEEPLX_CONNECTION_POOL = int(os.environ.get("DEEPLX_CONNECTION_POOL", str(DEEPLX_MAX_CONCURRENCY * 4)))

deeplx_session = requests.Session()
deeplx_adapter = HTTPAdapter(pool_connections=DEEPLX_CONNECTION_POOL, pool_maxsize=DEEPLX_CONNECTION_POOL)
deeplx_session.mount("http://", deeplx_adapter)
deeplx_session.mount("https://", deeplx_adapter)

_DEEPLX_HEALTH_CACHE = {"ts": 0.0, "ok": False}

# 文件清理配置（秒）
FILE_MAX_AGE = 24 * 60 * 60  # 24小时


def cleanup_old_files():
    """清理超过24小时的上传文件和输出文件"""
    current_time = time.time()
    cleaned_count = 0

    # 清理uploads目录
    if os.path.exists(UPLOAD_FOLDER):
        for filename in os.listdir(UPLOAD_FOLDER):
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            if os.path.isfile(filepath):
                file_age = current_time - os.path.getmtime(filepath)
                if file_age > FILE_MAX_AGE:
                    try:
                        os.remove(filepath)
                        cleaned_count += 1
                        print(f"✓ 清理旧文件: {filename}")
                    except Exception as e:
                        print(f"⚠️ 清理文件失败 {filename}: {e}")

    # 清理outputs目录
    if os.path.exists(OUTPUT_FOLDER):
        for filename in os.listdir(OUTPUT_FOLDER):
            filepath = os.path.join(OUTPUT_FOLDER, filename)
            if os.path.isfile(filepath):
                file_age = current_time - os.path.getmtime(filepath)
                if file_age > FILE_MAX_AGE:
                    try:
                        os.remove(filepath)
                        cleaned_count += 1
                        print(f"✓ 清理旧文件: {filename}")
                    except Exception as e:
                        print(f"⚠️ 清理文件失败 {filename}: {e}")

    if cleaned_count > 0:
        print(f"✅ 共清理 {cleaned_count} 个过期文件")


def start_cleanup_scheduler():
    """启动定时清理任务"""
    import time

    def cleanup_loop():
        while True:
            time.sleep(3600)  # 每小时检查一次
            cleanup_old_files()

    cleanup_thread = threading.Thread(target=cleanup_loop, daemon=True)
    cleanup_thread.start()
    print("🧹 文件清理服务已启动（每小时检查，清理超过24小时的文件）")


def _is_pdf_file(path: str) -> bool:
    try:
        with open(path, "rb") as f:
            return f.read(5) == b"%PDF-"
    except OSError:
        return False


class ParseTask:
    def __init__(self, task_id: str, filename: str):
        self.task_id = task_id
        self.filename = filename
        self.status = "pending"
        self.progress = 0
        self.message = "等待处理..."
        self.result = None
        self.error = None
        self.pdf_path = None
        self.created_at = datetime.now()

    def to_dict(self):
        return {
            "task_id": self.task_id,
            "filename": self.filename,
            "status": self.status,
            "progress": self.progress,
            "message": self.message,
            "error": self.error,
            "created_at": self.created_at.isoformat()
        }


def process_parse(task_id: str, input_path: str):
    task = tasks.get(task_id)
    if not task:
        return
    try:
        task.status = "parsing"
        task.progress = 10
        task.message = "正在上传文件..."
        parser = DocumentParser()
        task.progress = 30
        task.message = "正在解析文档..."
        result = parser.parse(file_path=input_path)
        task.progress = 90
        task.message = "处理完成"
        task.result = result
        # 仅在原文件为PDF时才提供预览
        if input_path.lower().endswith(".pdf"):
            task_output_dir = os.path.join(OUTPUT_FOLDER, task_id)
            os.makedirs(task_output_dir, exist_ok=True)
            pdf_dest = os.path.join(task_output_dir, "original.pdf")
            shutil.copy2(input_path, pdf_dest)
            task.pdf_path = pdf_dest
        task.status = "completed"
        task.progress = 100
        task.message = "解析完成!"
    except Exception as e:
        task.status = "failed"
        task.error = str(e)
        task.message = f"处理失败: {str(e)}"
        import traceback
        traceback.print_exc()
        # 解析失败时删除上传的文件
        if os.path.exists(input_path):
            try:
                os.remove(input_path)
                print(f"✓ 已删除上传文件: {input_path}")
            except Exception as e2:
                print(f"⚠️ 删除文件失败: {e2}")


def _deeplx_health_check(texts, source_lang, target_lang):
    """对DeepLX做缓存健康检测，避免每次批量调用都额外请求一次"""
    now = time.time()
    last_ts = _DEEPLX_HEALTH_CACHE.get("ts", 0.0)
    if (now - last_ts) < DEEPLX_HEALTH_TTL and _DEEPLX_HEALTH_CACHE.get("ok"):
        return

    sample = next((t for t in texts if (t or "").strip()), "")
    if not sample:
        return

    payload = {
        "text": sample[:200],
        "source_lang": source_lang or DEFAULT_SOURCE_LANG,
        "target_lang": target_lang or DEFAULT_TARGET_LANG,
    }

    try:
        deeplx_session.post(
            DEEPLX_API_URL,
            json=payload,
            timeout=(3, min(15, DEEPLX_TIMEOUT))
        )
        _DEEPLX_HEALTH_CACHE.update({"ts": now, "ok": True})
    except Exception as exc:
        _DEEPLX_HEALTH_CACHE.update({"ts": now, "ok": False})
        print(f"⚠️ DeepLX连通性检测失败（将继续尝试翻译）: {exc}", flush=True)


def translate_via_deeplx(text: str, source_lang: str, target_lang: str) -> str:
    if not DEEPLX_API_URL:
        raise ValueError("DEEPLX_API_URL 未配置")

    payload = {
        "text": text,
        "source_lang": source_lang or DEFAULT_SOURCE_LANG,
        "target_lang": target_lang or DEFAULT_TARGET_LANG
    }

    last_error = None
    for attempt in range(1, DEEPLX_MAX_RETRIES + 1):
        try:
            if attempt > 1:
                wait_time = min(2 ** (attempt - 1), 5)
                print(f"⏳ DeepLX重试等待 {wait_time} 秒（第 {attempt}/{DEEPLX_MAX_RETRIES} 次）...", flush=True)
                time.sleep(wait_time)

            start_ts = time.time()
            response = deeplx_session.post(
                DEEPLX_API_URL,
                json=payload,
                timeout=(5, DEEPLX_TIMEOUT)
            )
            elapsed = time.time() - start_ts
            print(f"🔁 DeepLX响应: status={response.status_code} time={elapsed:.2f}s", flush=True)

            if response.status_code == 429:
                raise ValueError("DeepLX接口触发限流(429)，请稍后重试或降低批量大小/频率。")

            response.raise_for_status()
            data = response.json()

            if data.get("code") == 200:
                return data.get("data", text)

            raise ValueError(data.get("message") or data.get("msg") or "DeepLX翻译失败")
        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as exc:
            last_error = exc
            print(f"⚠️ DeepLX请求异常: {exc}", flush=True)
            continue
        except Exception as exc:
            last_error = exc
            print(f"⚠️ DeepLX翻译错误: {exc}", flush=True)
            continue

    raise ValueError(f"DeepLX翻译失败: {last_error}")


def translate_via_openai(text: str, config: dict, source_lang: str, target_lang: str) -> str:
    base_url = (config or {}).get("base_url")
    api_key = (config or {}).get("api_key")
    model = (config or {}).get("model")

    if not base_url or not api_key or not model:
        raise ValueError("缺少OpenAI翻译配置")

    url = base_url.rstrip('/') + "/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    prompt = f"Translate the following text from {source_lang or 'Source Language'} to {target_lang or 'Target Language'}:\n\n{text}"
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "You are a professional translation engine."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.2,
        "stream": False
    }

    response = requests.post(url, headers=headers, json=payload, timeout=90)
    response.raise_for_status()
    data = response.json()
    choices = data.get("choices", [])
    if not choices:
        raise ValueError("OpenAI翻译返回为空")
    return choices[0]["message"]["content"].strip()


def batch_translate_via_deeplx(texts, source_lang, target_lang):
    if not texts:
        return []

    _deeplx_health_check(texts, source_lang, target_lang)

    # 优化：合并小文本块以减少请求次数
    MAX_CHUNK_SIZE = 3000  # DeepLX可以处理较大的文本块
    merged_chunks = []
    current_chunk = []
    current_size = 0

    for text in texts:
        if not text or not text.strip():
            # 保留空文本的位置
            if current_chunk:
                merged_chunks.append(("\n\n".join(current_chunk), len(current_chunk)))
                current_chunk = []
                current_size = 0
            merged_chunks.append(("", 1))
            continue

        text_len = len(text)
        # 如果单个文本就超过限制，单独处理
        if text_len > MAX_CHUNK_SIZE:
            if current_chunk:
                merged_chunks.append(("\n\n".join(current_chunk), len(current_chunk)))
                current_chunk = []
                current_size = 0
            merged_chunks.append((text, 1))
        # 如果加上当前文本会超过限制，先保存当前块
        elif current_size + text_len + 2 > MAX_CHUNK_SIZE:
            if current_chunk:
                merged_chunks.append(("\n\n".join(current_chunk), len(current_chunk)))
            current_chunk = [text]
            current_size = text_len
        # 否则累积到当前块
        else:
            current_chunk.append(text)
            current_size += text_len + 2

    # 保存最后一块
    if current_chunk:
        merged_chunks.append(("\n\n".join(current_chunk), len(current_chunk)))

    print(f"📊 优化前: {len(texts)} 个文本块，优化后: {len(merged_chunks)} 个请求", flush=True)

    # 使用线程池并发翻译
    results = []
    success_count = 0
    last_error = None

    def translate_chunk(chunk_text):
        if not chunk_text:
            return chunk_text
        try:
            return translate_via_deeplx(chunk_text, source_lang, target_lang)
        except Exception as exc:
            print(f"⚠️ 翻译块失败: {exc}", flush=True)
            return chunk_text

    with ThreadPoolExecutor(max_workers=DEEPLX_MAX_CONCURRENCY) as executor:
        futures = []
        for chunk_text, count in merged_chunks:
            future = executor.submit(translate_chunk, chunk_text)
            futures.append((future, count))
            # 添加小延迟避免瞬间大量请求
            if DEEPLX_RATE_LIMIT > 0:
                time.sleep(DEEPLX_RATE_LIMIT / DEEPLX_MAX_CONCURRENCY)

        # 收集结果
        for future, count in futures:
            try:
                translated = future.result()
                if not translated:
                    results.append(translated)
                elif count == 1:
                    results.append(translated)
                    if translated:
                        success_count += 1
                else:
                    # 拆分合并的块
                    parts = translated.split("\n\n")
                    results.extend(parts)
                    success_count += len([p for p in parts if p])
            except Exception as exc:
                last_error = exc
                results.extend([""] * count)

    if success_count == 0 and last_error is not None:
        raise ValueError(str(last_error))

    return results


def batch_translate_via_openai(texts, config, source_lang, target_lang):
    results = []
    for text in texts:
        if not text:
            results.append(text)
            continue
        translated = translate_via_openai(text, config, source_lang, target_lang)
        results.append(translated)
    return results


@app.route('/')
def serve_frontend_index():
    if FRONTEND_DIST.joinpath('index.html').exists():
        return send_from_directory(FRONTEND_DIST, 'index.html')
    return render_template('index.html')


@app.route('/viewer')
def render_viewer():
    return render_template('index.html')


@app.route('/assets/<path:filename>')
def serve_frontend_assets(filename):
    assets_dir = FRONTEND_DIST / 'assets'
    if assets_dir.exists():
        return send_from_directory(assets_dir, filename)
    return jsonify({"error": "前端静态资源尚未构建"}), 404


@app.route('/<path:path>')
def serve_frontend_static(path):
    target = FRONTEND_DIST / path
    if target.exists() and target.is_file():
        return send_from_directory(FRONTEND_DIST, path)
    return serve_frontend_index()


@app.route('/api/parse', methods=['POST'])
def parse():
    if 'file' not in request.files:
        return jsonify({"error": "没有上传文件"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "文件名为空"}), 400
    task_id = str(uuid.uuid4())[:8]
    filename = f"{task_id}_{file.filename}"
    input_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(input_path)
    task = ParseTask(task_id, file.filename)
    tasks[task_id] = task
    thread = threading.Thread(target=process_parse, args=(task_id, input_path))
    thread.daemon = True
    thread.start()
    return jsonify({"task_id": task_id, "message": "任务已创建"})


@app.route('/api/task/<task_id>')
def get_task(task_id):
    task = tasks.get(task_id)
    if not task:
        return jsonify({"error": "任务不存在"}), 404
    return jsonify(task.to_dict())


@app.route('/api/result/<task_id>')
def get_result(task_id):
    task = tasks.get(task_id)
    if not task:
        return jsonify({"error": "任务不存在"}), 404
    if task.status != 'completed' or not task.result:
        return jsonify({"error": "结果未准备好"}), 400
    return jsonify({
        "markdown": task.result.get("markdown", ""),
        "content_list": task.result.get("content_list", []),
        "page_mappings": task.result.get("page_mappings", {})
    })


@app.route('/api/pdf/<task_id>')
def get_pdf(task_id):
    # 优先从任务对象获取
    task = tasks.get(task_id)
    if task and task.pdf_path and os.path.exists(task.pdf_path):
        if _is_pdf_file(task.pdf_path):
            print(f"✓ 返回PDF文件(从任务): {task.pdf_path}")
            return send_file(task.pdf_path, mimetype='application/pdf')
        print(f"⚠️ PDF请求失败: 文件不是PDF {task.pdf_path}")
    # 任务对象不存在时，尝试从outputs目录查找
    pdf_path = os.path.join(OUTPUT_FOLDER, task_id, "original.pdf")
    if os.path.exists(pdf_path) and _is_pdf_file(pdf_path):
        print(f"✓ 返回PDF文件(从outputs): {pdf_path}")
        return send_file(pdf_path, mimetype='application/pdf')
    print(f"⚠️ PDF请求失败: 任务 {task_id} 的PDF文件不存在")
    return jsonify({"error": "PDF文件不存在"}), 404


@app.route('/api/image/<task_id>/<image_name>')
def get_image(task_id, image_name):
    task = tasks.get(task_id)
    if not task or not task.result:
        return jsonify({"error": "任务不存在"}), 404
    images = task.result.get("images", {})
    if image_name not in images:
        return jsonify({"error": "图片不存在"}), 404
    return send_file(BytesIO(images[image_name]), mimetype='image/png')


@app.route('/api/translate', methods=['POST'])
def translate_text():
    data = request.get_json(force=True)
    text = (data.get("text") or "").strip()
    if not text:
        return jsonify({"error": "缺少文本内容"}), 400

    provider_raw = data.get("provider", "deeplx")
    provider = (provider_raw or "").strip().lower()
    source_lang = data.get("source_lang") or DEFAULT_SOURCE_LANG
    target_lang = data.get("target_lang") or DEFAULT_TARGET_LANG

    try:
        if provider == "openai":
            translated = translate_via_openai(text, data.get("config"), source_lang, target_lang)
        else:
            translated = translate_via_deeplx(text, source_lang, target_lang)
        return jsonify({"translated_text": translated})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/translate_batch', methods=['POST'])
def translate_batch():
    data = request.get_json(force=True)
    chunks = data.get("chunks")
    if not isinstance(chunks, list) or not chunks:
        return jsonify({"error": "缺少翻译内容"}), 400

    texts = [(chunk.get("text") or "") for chunk in chunks]
    provider_raw = data.get("provider", "deeplx")
    provider = (provider_raw or "").strip().lower()
    source_lang = data.get("source_lang") or DEFAULT_SOURCE_LANG
    target_lang = data.get("target_lang") or DEFAULT_TARGET_LANG
    print(
        f"🌍 translate_batch provider_raw={provider_raw!r} provider={provider} chunks={len(texts)} source={source_lang} target={target_lang}",
        flush=True
    )
    if texts:
        lengths = [len(t or "") for t in texts]
        max_len = max(lengths) if lengths else 0
        min_len = min(lengths) if lengths else 0
        print(f"🧩 chunk_len min={min_len} max={max_len}", flush=True)

    try:
        if provider == "openai":
            print("➡️ using openai", flush=True)
            translations = batch_translate_via_openai(texts, data.get("config"), source_lang, target_lang)
        else:
            print("➡️ using deeplx", flush=True)
            translations = batch_translate_via_deeplx(texts, source_lang, target_lang)
        return jsonify({"translations": translations})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/download_pdf/<task_id>', methods=['POST'])
def download_translated_pdf(task_id):
    """下载翻译后的PDF文件"""
    task = tasks.get(task_id)
    if not task:
        return jsonify({"error": "任务不存在"}), 404
    if task.status != 'completed' or not task.result:
        return jsonify({"error": "任务未完成"}), 400

    data = request.get_json(force=True) if request.is_json else {}
    content_list = data.get("content_list", [])

    try:
        from pdf_generator import markdown_to_pdf, content_list_to_markdown

        # 获取图片的辅助函数
        def get_image_data(tid, image_name):
            t = tasks.get(tid)
            if t and t.result:
                imgs = t.result.get("images", {})
                return imgs.get(image_name)
            return None

        # 构建Markdown内容
        if content_list:
            # 使用前端传来的翻译后内容
            markdown_content = content_list_to_markdown(content_list, task.result.get("images", {}))
        else:
            # 使用原始Markdown
            markdown_content = task.result.get("markdown", "")

        if not markdown_content:
            return jsonify({"error": "没有可导出的内容"}), 400

        # 生成PDF
        output_filename = f"translated_{task_id}.pdf"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)

        markdown_to_pdf(
            markdown_content,
            output_path,
            images=task.result.get("images", {}),
            task_id=task_id,
            get_image_func=get_image_data
        )

        # 在响应完成后删除生成的PDF文件
        @after_this_request
        def cleanup(response):
            try:
                if os.path.exists(output_path):
                    os.remove(output_path)
                    print(f"✓ 已删除临时PDF: {output_path}")
            except Exception as e:
                print(f"⚠️ 删除临时PDF失败: {e}")
            return response

        return send_file(
            output_path,
            as_attachment=True,
            download_name=f"translated_{task.filename}",
            mimetype='application/pdf'
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"PDF生成失败: {str(e)}"}), 500


@app.route('/api/download/<task_id>/<format_type>', methods=['POST'])
def download_document(task_id, format_type):
    """下载多种格式的文档"""
    task = tasks.get(task_id)
    if not task:
        return jsonify({"error": "任务不存在"}), 404
    if task.status != 'completed' or not task.result:
        return jsonify({"error": "任务未完成"}), 400

    data = request.get_json(force=True) if request.is_json else {}
    content_list = data.get("content_list", [])

    # 获取基础文件名（去掉扩展名）
    base_filename = os.path.splitext(task.filename)[0]

    try:
        if format_type == 'markdown':
            return download_as_markdown(task, content_list, base_filename)
        elif format_type == 'html':
            return download_as_html(task, content_list, base_filename)
        elif format_type == 'docx':
            return download_as_docx(task, content_list, base_filename)
        elif format_type == 'json':
            return download_as_json(task, content_list, base_filename)
        elif format_type == 'latex':
            return download_as_latex(task, content_list, base_filename)
        else:
            return jsonify({"error": f"不支持的格式: {format_type}"}), 400
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"导出失败: {str(e)}"}), 500


def get_content_markdown(task, content_list):
    """从content_list或task.result获取Markdown内容"""
    from pdf_generator import content_list_to_markdown
    if content_list:
        return content_list_to_markdown(content_list, task.result.get("images", {}))
    return task.result.get("markdown", "")


def download_as_markdown(task, content_list, base_filename):
    """下载为Markdown格式"""
    from pdf_generator import content_list_to_markdown

    if content_list:
        markdown_content = content_list_to_markdown(content_list, task.result.get("images", {}))
    else:
        markdown_content = task.result.get("markdown", "")

    if not markdown_content:
        return jsonify({"error": "没有可导出的内容"}), 400

    # 创建临时文件
    output_path = os.path.join(OUTPUT_FOLDER, f"{task.task_id}_export.md")
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown_content)

    @after_this_request
    def cleanup(response):
        try:
            if os.path.exists(output_path):
                os.remove(output_path)
        except Exception:
            pass
        return response

    return send_file(
        output_path,
        as_attachment=True,
        download_name=f"{base_filename}.md",
        mimetype='text/markdown'
    )


def _has_translated_content(content_list):
    """检查content_list中是否包含翻译后的内容"""
    if not content_list:
        return False
    for item in content_list:
        if item.get('translated_text'):
            return True
    return False


def download_as_html(task, content_list, base_filename):
    """下载为HTML格式 - 有翻译内容时使用本地转换，否则用MinerU原文"""
    # 检查是否有翻译内容
    has_translation = _has_translated_content(content_list)

    # 如果没有翻译内容，可以使用MinerU生成的原文HTML
    if not has_translation:
        export_files = task.result.get("export_files", {})
        mineru_html = export_files.get("html")
        if mineru_html:
            print(f"✓ 使用MinerU生成的HTML文件（原文）")
            output_path = os.path.join(OUTPUT_FOLDER, f"{task.task_id}_export.html")
            with open(output_path, 'wb') as f:
                f.write(mineru_html)

            @after_this_request
            def cleanup(response):
                try:
                    if os.path.exists(output_path):
                        os.remove(output_path)
                except Exception:
                    pass
                return response

            return send_file(
                output_path,
                as_attachment=True,
                download_name=f"{base_filename}.html",
                mimetype='text/html'
            )

    # 有翻译内容，使用本地转换
    print(f"✓ 使用本地转换生成HTML（含译文）")
    import markdown
    from pdf_generator import content_list_to_markdown

    if content_list:
        markdown_content = content_list_to_markdown(content_list, task.result.get("images", {}))
    else:
        markdown_content = task.result.get("markdown", "")

    if not markdown_content:
        return jsonify({"error": "没有可导出的内容"}), 400

    # 转换Markdown为HTML
    md = markdown.Markdown(extensions=['tables', 'fenced_code', 'toc'])
    html_body = md.convert(markdown_content)

    # 处理图片 - 将图片转为base64内嵌
    images = task.result.get("images", {})
    for img_name, img_data in images.items():
        if img_data:
            img_base64 = base64.b64encode(img_data).decode('utf-8')
            html_body = html_body.replace(
                f'/api/image/{task.task_id}/{img_name}',
                f'data:image/png;base64,{img_base64}'
            )
            html_body = html_body.replace(
                f'images/{img_name}',
                f'data:image/png;base64,{img_base64}'
            )

    html_content = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{base_filename}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 900px; margin: 0 auto; padding: 40px 20px; line-height: 1.8; color: #333; }}
        h1, h2, h3 {{ margin-top: 1.5em; color: #1a1a1a; }}
        img {{ max-width: 100%; height: auto; }}
        pre {{ background: #f5f5f5; padding: 15px; border-radius: 5px; overflow-x: auto; }}
        code {{ background: #f5f5f5; padding: 2px 6px; border-radius: 3px; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px 12px; }}
    </style>
</head>
<body>
{html_body}
</body>
</html>'''

    output_path = os.path.join(OUTPUT_FOLDER, f"{task.task_id}_export.html")
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_content)

    @after_this_request
    def cleanup(response):
        try:
            if os.path.exists(output_path):
                os.remove(output_path)
        except Exception:
            pass
        return response

    return send_file(
        output_path,
        as_attachment=True,
        download_name=f"{base_filename}.html",
        mimetype='text/html'
    )


def download_as_docx(task, content_list, base_filename):
    """下载为DOCX格式 - 有翻译内容时使用本地转换，否则用MinerU原文"""
    # 检查是否有翻译内容
    has_translation = _has_translated_content(content_list)

    # 如果没有翻译内容，可以使用MinerU生成的原文DOCX
    if not has_translation:
        export_files = task.result.get("export_files", {})
        mineru_docx = export_files.get("docx")
        if mineru_docx:
            print(f"✓ 使用MinerU生成的DOCX文件（原文）")
            output_path = os.path.join(OUTPUT_FOLDER, f"{task.task_id}_export.docx")
            with open(output_path, 'wb') as f:
                f.write(mineru_docx)

            @after_this_request
            def cleanup(response):
                try:
                    if os.path.exists(output_path):
                        os.remove(output_path)
                except Exception:
                    pass
                return response

            return send_file(
                output_path,
                as_attachment=True,
                download_name=f"{base_filename}.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

    # 有翻译内容，使用本地转换
    print(f"✓ 使用本地转换生成DOCX（含译文）")
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document()
    title = doc.add_heading(base_filename, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if content_list:
        for item in content_list:
            item_type = item.get('type', 'text')
            text = item.get('translated_text') or item.get('text', '')
            text_level = item.get('text_level')

            if item_type == 'image':
                img_path = item.get('img_path', '')
                img_name = img_path.split('/')[-1] if img_path else ''
                images = task.result.get("images", {})
                if img_name and img_name in images:
                    img_data = images[img_name]
                    if img_data:
                        try:
                            img_stream = BytesIO(img_data)
                            doc.add_picture(img_stream, width=Inches(5))
                        except Exception:
                            pass
                caption = item.get('image_caption', [])
                if caption:
                    p = doc.add_paragraph(' '.join(caption))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif item_type == 'list':
                if text:
                    doc.add_paragraph(text)
                else:
                    list_items = item.get('list_items', [])
                    for li in list_items:
                        doc.add_paragraph(li, style='List Bullet')
            elif text_level and text_level <= 6:
                doc.add_heading(text, level=min(text_level, 9))
            elif text:
                doc.add_paragraph(text)
    else:
        markdown_content = task.result.get("markdown", "")
        for line in markdown_content.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            else:
                doc.add_paragraph(line)

    output_path = os.path.join(OUTPUT_FOLDER, f"{task.task_id}_export.docx")
    doc.save(output_path)

    @after_this_request
    def cleanup(response):
        try:
            if os.path.exists(output_path):
                os.remove(output_path)
        except Exception:
            pass
        return response

    return send_file(
        output_path,
        as_attachment=True,
        download_name=f"{base_filename}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


def download_as_json(task, content_list, base_filename):
    """下载为JSON格式（包含完整结构信息）"""
    # 构建完整的JSON数据
    export_data = {
        "filename": task.filename,
        "task_id": task.task_id,
        "created_at": task.created_at.isoformat(),
        "content_list": content_list if content_list else task.result.get("content_list", []),
        "markdown": task.result.get("markdown", ""),
        "page_mappings": task.result.get("page_mappings", {}),
        "metadata": {
            "export_time": datetime.now().isoformat(),
            "format_version": "1.0"
        }
    }

    # 不导出图片的二进制数据，只导出图片名称列表
    images = task.result.get("images", {})
    export_data["image_names"] = list(images.keys())

    output_path = os.path.join(OUTPUT_FOLDER, f"{task.task_id}_export.json")
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(export_data, f, ensure_ascii=False, indent=2)

    @after_this_request
    def cleanup(response):
        try:
            if os.path.exists(output_path):
                os.remove(output_path)
        except Exception:
            pass
        return response

    return send_file(
        output_path,
        as_attachment=True,
        download_name=f"{base_filename}.json",
        mimetype='application/json'
    )


def download_as_latex(task, content_list, base_filename):
    """下载为LaTeX格式 - 有翻译内容时使用本地转换，否则用MinerU原文"""
    # 检查是否有翻译内容
    has_translation = _has_translated_content(content_list)

    # 如果没有翻译内容，可以使用MinerU生成的原文LaTeX
    if not has_translation:
        export_files = task.result.get("export_files", {})
        mineru_latex = export_files.get("latex")
        if mineru_latex:
            print(f"✓ 使用MinerU生成的LaTeX文件（原文）")
            output_path = os.path.join(OUTPUT_FOLDER, f"{task.task_id}_export.tex")
            with open(output_path, 'wb') as f:
                f.write(mineru_latex)

            @after_this_request
            def cleanup(response):
                try:
                    if os.path.exists(output_path):
                        os.remove(output_path)
                except Exception:
                    pass
                return response

            return send_file(
                output_path,
                as_attachment=True,
                download_name=f"{base_filename}.tex",
                mimetype='application/x-tex'
            )

    # 有翻译内容，使用本地转换
    print(f"✓ 使用本地转换生成LaTeX（含译文）")
    from pdf_generator import content_list_to_markdown
    import re

    if content_list:
        markdown_content = content_list_to_markdown(content_list, task.result.get("images", {}))
    else:
        markdown_content = task.result.get("markdown", "")

    if not markdown_content:
        return jsonify({"error": "没有可导出的内容"}), 400

    latex_content = markdown_to_latex(markdown_content, base_filename)

    output_path = os.path.join(OUTPUT_FOLDER, f"{task.task_id}_export.tex")
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(latex_content)

    @after_this_request
    def cleanup(response):
        try:
            if os.path.exists(output_path):
                os.remove(output_path)
        except Exception:
            pass
        return response

    return send_file(
        output_path,
        as_attachment=True,
        download_name=f"{base_filename}.tex",
        mimetype='application/x-tex'
    )


def markdown_to_latex(markdown_content, title="Document"):
    """将Markdown转换为LaTeX格式"""
    import re

    # LaTeX文档头
    latex = r'''\documentclass[12pt,a4paper]{article}
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage{xeCJK}
\usepackage{graphicx}
\usepackage{hyperref}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{listings}
\usepackage{xcolor}
\usepackage{geometry}
\geometry{margin=2.5cm}

\lstset{
    basicstyle=\ttfamily\small,
    breaklines=true,
    frame=single,
    backgroundcolor=\color{gray!10}
}

\title{''' + escape_latex(title) + r'''}
\date{\today}

\begin{document}
\maketitle

'''

    lines = markdown_content.split('\n')
    in_code_block = False
    code_lang = ''

    for line in lines:
        # 代码块处理
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line.strip()[3:]
                latex += '\\begin{lstlisting}'
                if code_lang:
                    latex += f'[language={code_lang}]'
                latex += '\n'
            else:
                in_code_block = False
                latex += '\\end{lstlisting}\n'
            continue

        if in_code_block:
            latex += line + '\n'
            continue

        # 标题处理
        if line.startswith('# '):
            latex += '\\section{' + escape_latex(line[2:]) + '}\n'
        elif line.startswith('## '):
            latex += '\\subsection{' + escape_latex(line[3:]) + '}\n'
        elif line.startswith('### '):
            latex += '\\subsubsection{' + escape_latex(line[4:]) + '}\n'
        elif line.startswith('#### '):
            latex += '\\paragraph{' + escape_latex(line[5:]) + '}\n'
        # 列表处理
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            latex += '\\begin{itemize}\n'
            latex += '\\item ' + escape_latex(line.strip()[2:]) + '\n'
            latex += '\\end{itemize}\n'
        elif re.match(r'^\d+\. ', line.strip()):
            latex += '\\begin{enumerate}\n'
            latex += '\\item ' + escape_latex(re.sub(r'^\d+\. ', '', line.strip())) + '\n'
            latex += '\\end{enumerate}\n'
        # 图片处理
        elif '![' in line:
            match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if match:
                alt_text = match.group(1)
                img_path = match.group(2)
                latex += '\\begin{figure}[h]\n'
                latex += '\\centering\n'
                latex += f'% \\includegraphics[width=0.8\\textwidth]{{{img_path}}}\n'
                if alt_text:
                    latex += f'\\caption{{{escape_latex(alt_text)}}}\n'
                latex += '\\end{figure}\n'
        # 数学公式处理 - 保持原样
        elif line.strip().startswith('$$') or line.strip().endswith('$$'):
            latex += line + '\n'
        elif '$' in line:
            latex += process_inline_math(line) + '\n'
        # 空行
        elif not line.strip():
            latex += '\n'
        # 普通段落
        else:
            latex += escape_latex(line) + '\n'

    latex += r'''
\end{document}
'''
    return latex


def escape_latex(text):
    """转义LaTeX特殊字符"""
    if not text:
        return ''
    # 保留数学公式中的内容
    parts = []
    last_end = 0
    # 匹配 $...$ 或 $$...$$ 的数学公式
    import re
    for match in re.finditer(r'\$\$.*?\$\$|\$.*?\$', text):
        # 转义公式前的文本
        before = text[last_end:match.start()]
        before = _escape_latex_chars(before)
        parts.append(before)
        # 保持公式原样
        parts.append(match.group())
        last_end = match.end()
    # 转义剩余文本
    after = text[last_end:]
    after = _escape_latex_chars(after)
    parts.append(after)
    return ''.join(parts)


def _escape_latex_chars(text):
    """转义LaTeX特殊字符（不含数学公式）"""
    if not text:
        return ''
    chars = {
        '&': r'\&',
        '%': r'\%',
        '#': r'\#',
        '_': r'\_',
        '{': r'\{',
        '}': r'\}',
        '~': r'\textasciitilde{}',
        '^': r'\textasciicircum{}',
    }
    for char, replacement in chars.items():
        text = text.replace(char, replacement)
    # 处理反斜杠（但不处理已有的LaTeX命令）
    text = re.sub(r'\\(?![a-zA-Z])', r'\\textbackslash{}', text)
    return text


def process_inline_math(line):
    """处理行内数学公式"""
    # 保持 $...$ 格式的公式不变
    return escape_latex(line)


# 启动时执行一次清理并启动定时清理服务
cleanup_old_files()
start_cleanup_scheduler()


if __name__ == '__main__':
    print("=" * 50)
    print("📄 PDF文档解析查看器")
    print("=" * 50)
    print("🌐 访问地址: http://localhost:8080")
    print("=" * 50)
    app.run(host='0.0.0.0', port=8080, debug=True)
